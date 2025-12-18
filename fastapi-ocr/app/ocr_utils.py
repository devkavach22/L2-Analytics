import os
import io
import fitz               # PyMuPDF
import easyocr
from PIL import Image
from docx import Document
import pandas as pd
import docx2txt

# Load EasyOCR reader once
READER = easyocr.Reader(['en'], gpu=False)

# ---------------------------------------------
# OCR image
# ---------------------------------------------
def ocr_image(image_bytes: bytes) -> str:
    import numpy as np
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    arr = np.array(img)
    results = READER.readtext(arr, detail=0)
    return "\n".join(results).strip()

# ---------------------------------------------
# PDF extraction (OCR + embedded text)
# ---------------------------------------------
def extract_pdf(pdf_bytes: bytes) -> str:
    text_pages = []
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page_index, page in enumerate(pdf):
        # 1️⃣ Extract embedded text (if real text exists)
        embedded_text = page.get_text().strip()

        # 2️⃣ If page has little/no embedded text → fallback to OCR
        if len(embedded_text) < 20:
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            ocr_text = ocr_image(img_bytes)
            text_pages.append(f"--- PAGE {page_index+1} ---\n{ocr_text}")
        else:
            text_pages.append(f"--- PAGE {page_index+1} ---\n{embedded_text}")

    return "\n\n".join(text_pages).strip()

# ---------------------------------------------
# Extract DOCX including tables
# ---------------------------------------------
def extract_docx(docx_bytes: bytes) -> str:
    try:
        file_like = io.BytesIO(docx_bytes)
        doc = Document(file_like)

        lines = []
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        # Extract tables (ATS Friendly Resume tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    lines.append(" | ".join(row_text))

        return "\n".join(lines).strip()
    except:
        # More accurate DOCX extraction fallback
        try:
            return docx2txt.process(io.BytesIO(docx_bytes))
        except:
            return ""

# ---------------------------------------------
# Read simple text formats
# ---------------------------------------------
def read_simple_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except:
        return file_bytes.decode("latin-1", errors="ignore")

# ---------------------------------------------
# XLSX extraction to readable text
# ---------------------------------------------
def extract_xlsx(xlsx_bytes: bytes) -> str:
    try:
        file_like = io.BytesIO(xlsx_bytes)
        df = pd.read_excel(file_like)
        return df.to_string(index=False)
    except:
        return ""

# ---------------------------------------------
# Universal text extractor
# ---------------------------------------------
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    # Images
    if ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"]:
        return ocr_image(file_bytes)

    # PDF (text + scanned resume OCR)
    if ext == ".pdf":
        return extract_pdf(file_bytes)

    # DOCX (ATS-friendly resumes)
    if ext == ".docx":
        return extract_docx(file_bytes)

    # DOC (convert-like fallback)
    if ext == ".doc":
        return "DOC extraction not fully supported. Convert to DOCX for best results."

    # Excel
    if ext == ".xlsx":
        return extract_xlsx(file_bytes)

    # Simple text-based
    if ext in [".txt", ".csv", ".json", ".md", ".log"]:
        return read_simple_text(file_bytes)

    # Unknown → attempt OCR
    try:
        return ocr_image(file_bytes)
    except:
        return ""
